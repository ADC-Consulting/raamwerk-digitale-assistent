#!/usr/bin/env python3
"""
Generate a per-good-practice .docx review document.

Page 1: practice → domains overview matrix so reviewers can scan multi-domain
        membership at a glance.
Per-practice pages (one per practice, with a page break between):
  - Title, summary
  - phases, levels
  - domains (resolved to titles)
  - body (recommendations) as paragraphs
  - all sources resolved from bronnen.yaml: title, categorie, omschrijving, url

Output: local_output/review/good-practices-review.docx
"""

import os
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

from _content import (
    REPO_ROOT,
    load_bronnen,
    load_domain_files,
    load_practice_files,
)

OUT_DIR = os.path.join(REPO_ROOT, 'local_output', 'review')
OUT_FILE = os.path.join(OUT_DIR, 'good-practices-review.docx')


def _set_run_color(run, hex_color):
    run.font.color.rgb = RGBColor.from_string(hex_color)


def write_overview(doc, practices, domains_by_id):
    """Page 1: practice → domains matrix."""
    h = doc.add_heading('Good practices — overzicht en domein-toewijzing', level=0)
    p = doc.add_paragraph(
        'Een overzicht van alle good practices, hun samenvatting, en welke '
        'domein(en) zij raken. Doel: snel beoordelen of de cross-domain-toewijzing '
        'past, voordat je per practice doorklikt.'
    )
    p.runs[0].font.size = Pt(10)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Practice'
    hdr[1].text = 'Domeinen'
    hdr[2].text = 'Phases / Levels'
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for p in sorted(practices, key=lambda x: x.get('title', '')):
        row = table.add_row().cells
        row[0].text = p.get('title', '')
        domain_titles = []
        for did in p.get('domains') or []:
            d = domains_by_id.get(did)
            domain_titles.append(d.get('title', did) if d else did)
        row[1].text = ', '.join(domain_titles)
        meta = []
        if p.get('phases'):
            meta.append('Phases: ' + ', '.join(p['phases']))
        if p.get('levels'):
            meta.append('Levels: ' + ', '.join(p['levels']))
        row[2].text = '\n'.join(meta)

    doc.add_paragraph()  # spacing before the page break
    doc.add_page_break()


def write_practice_page(doc, p, domains_by_id, bronnen_by_id):
    """One page per practice with title, meta, body, sources."""
    doc.add_heading(p.get('title', p.get('id', '')), level=1)

    # Practice id (small grey)
    pid = doc.add_paragraph()
    r = pid.add_run(f'id: {p.get("id", "")}')
    r.font.size = Pt(9)
    _set_run_color(r, '6B7785')

    # Summary
    summary = (p.get('summary') or '').strip()
    if summary:
        para = doc.add_paragraph(summary)
        para.runs[0].italic = True

    # Meta line
    meta_para = doc.add_paragraph()
    meta_para.add_run('Phases: ').bold = True
    meta_para.add_run(', '.join(p.get('phases') or []) + '    ')
    meta_para.add_run('Levels: ').bold = True
    meta_para.add_run(', '.join(p.get('levels') or []))

    # Domains (resolved to titles)
    domain_titles = []
    for did in p.get('domains') or []:
        d = domains_by_id.get(did)
        domain_titles.append(d.get('title', did) if d else did)
    dom_para = doc.add_paragraph()
    dom_para.add_run('Domeinen: ').bold = True
    dom_para.add_run(', '.join(domain_titles) or '—')

    # Body (Praktische tips)
    doc.add_heading('Praktische tips', level=2)
    for item in p.get('body') or []:
        doc.add_paragraph(item)

    # Sources
    sources = p.get('sources') or []
    if sources:
        doc.add_heading('Bronnen', level=2)
        for sid in sources:
            b = bronnen_by_id.get(sid) if isinstance(sid, str) else None
            if not b:
                doc.add_paragraph(f'(ongebonden bron-id: {sid})')
                continue
            sp = doc.add_paragraph()
            sp.add_run(b.get('title', sid)).bold = True
            sp.add_run('  —  ').italic = True
            cat = b.get('categorie') or ''
            if cat:
                sp.add_run(cat).italic = True
            oms = (b.get('omschrijving') or '').strip()
            if oms:
                desc_p = doc.add_paragraph(oms)
                desc_p.paragraph_format.left_indent = Pt(18)
                for run in desc_p.runs:
                    run.font.size = Pt(10)
            url = b.get('url') or ''
            if url:
                url_p = doc.add_paragraph()
                url_p.paragraph_format.left_indent = Pt(18)
                ru = url_p.add_run(url)
                ru.font.size = Pt(9)
                _set_run_color(ru, '01689B')

    # Page break (skipped after the final practice — we don't know in this fn,
    # but trailing page break is harmless)
    last_para = doc.paragraphs[-1]
    last_para.add_run().add_break(WD_BREAK.PAGE)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    practices = load_practice_files()
    domains = load_domain_files()
    bronnen = load_bronnen()
    domains_by_id = {d['_slug']: d for d in domains}
    bronnen_by_id = {b['id']: b for b in bronnen if 'id' in b}

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    write_overview(doc, practices, domains_by_id)

    for p in sorted(practices, key=lambda x: x.get('title', '')):
        write_practice_page(p, doc=doc, domains_by_id=domains_by_id, bronnen_by_id=bronnen_by_id) if False else write_practice_page(doc, p, domains_by_id, bronnen_by_id)

    doc.save(OUT_FILE)
    print(f'Wrote {OUT_FILE}')
    print(f'  practices: {len(practices)}')
    print(f'  domains:   {len(domains)}')
    print(f'  bronnen:   {len(bronnen)}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
